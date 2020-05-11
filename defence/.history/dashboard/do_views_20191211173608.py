from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.contrib.auth.models import User,Group
from .forms import UserCreationForm,TAapplicationForm,cemilacUserForm,proforma_A_form,commentsUploadForm
from django.contrib import messages
from common.decorators import role_required
from authmgmt.models import registration
from .models import TAapplicationmodel,proforma_A_model,TAapplicationfiles,statusmodel,commentsmodel,idgenerationmodel
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.http import HttpResponse
from .views import link_callback
import os
from os import stat, remove
import pyAesCrypt
from datetime import datetime
from django.utils import formats
import comtypes.client
import pythoncom
import urllib
from docx import Document
import io
from io import BytesIO,StringIO
# import io.StringIO
from django.core.files import File

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer"])
def process_proforma(request):
    reg=TAapplicationmodel.objects.filter(file_in_id=str(request.user.id))
    return render(request, 'dealing officer/viewtyperecord.html',{'details':reg,'status':True})

# @login_required(login_url=settings.LOGIN_URL)
# @role_required(allowed_roles=["TA Coordinator"])
# def checklist(request):
#     reg=TAapplicationmodel.objects.all()
#     return render(request, 'dealing officer/viewtyperecord.html',{'details':reg,'status':True})

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["TA Applicant","Dealing Officer","TA Coordinator","RD","TCS-GD","TCS-CE","TCS-Dealing Officer","TCS-TA Coordinator"])
def viewtyperecord(request,id):
    print('saiiiiiiiiiiiiiii',id)
    # reg=get_object_or_404(registration,id=id)
    # taa=TAapplicationmodel.objects.filter(user_id=id).first()
    # if request.method == 'POST':
    #     return render(request, 'dealing officer/newtypeapproval.html', {'form': form,})
    # else:
    #     form = TAapplicationForm(instance=taa)
    #     template = get_template('applicant/newtypeapprovalpdf.html')
    #     context= {
    #         'firmname':taa.firmname,
    #         'addr1':taa.addr1,
    #         'addr2':taa.addr2,
    #         'tot':taa.tot,
    #         'item_name':taa.item_name,
    #         'part_no':taa.part_no,
    #         'desc':taa.desc,
    #         'spec': taa.spec,
    #         'dal_mdi':taa.dal_mdi,
    #         'bom':taa.bom,
    #         'sop_acbs':taa.sop_acbs,
    #         'pc': taa.pc,
    #         'tre':taa.tre,
    #         'otheritems':taa.otheritems
    #     }
    #     response = HttpResponse(content_type='application/pdf')
    #     response['Content-Disposition'] = 'attachment; filename="report.pdf"'
    #     html = template.render(context)

    #     pisaStatus = pisa.CreatePDF(
    #     html, dest=response, link_callback=link_callback)
    #     if pisaStatus:
    #         return HttpResponse(response, content_type='application/pdf')
    # # if error then show some funy view
    #     if pisaStatus.err:
    #         return HttpResponse('We had some errors <pre>' + html + '</pre>')
    #     return response
    # return render(request, 'applicant/newtypeapprovalpdf.html', {'form': form,})

    # curr_path=curr_path.replace('/','\\')
    # new_path = os.path.join(settings.MEDIA_ROOT + curr_path)


    # with open(new_path+'TAapplication.pdf', 'rb') as pdf:
    #     response = HttpResponse(pdf.read(),content_type='application/pdf')
    #     response['Content-Disposition'] = 'filename=some_file.pdf'
    # return response
    print(id,'kkk')
    idprefix=request.POST['idprefix']
    filename=request.POST['filename']
    if filename!='':
        comment=request.POST['comment']
        if filename=="TAapplication.pdf":
            tf=TAapplicationfiles.objects.filter(user_id=id,filecategory="TAapplication").first()
            tf.comments=comment
            tf.save()
        pro=proforma_A_model.objects.all()
        messages.success(request, 'Comments Successfully Submitted !')  
    fc=TAapplicationmodel.objects.filter(user_id=id,idprefix=idprefix).first()
    print(fc.idprefix,'kkk')
    tafil=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory="TAapplication",refid=fc.idprefix).first()
    curr_path = "/"+str(fc.user_id)+"/"+fc.idprefix+"Annexure 1/TAapplication/"
    print(tafil,'tafile')
    filename='TAapplication.pdf'
    url='http://127.0.0.1:8000/media'+urllib.parse.quote(curr_path)+'TAapplication.pdf'
    print(tafil.comments,'new')
    return render(request, 'dealing officer/pdf viewer.html',{'url':url,'id':id,'filename':filename,'fc':tafil.comments,'idprefix':fc.idprefix})

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD","TCS-CE","TCS-Dealing Officer","TCS-TA Coordinator"])
def draft_ta(request,id):

    doc_final_path ='E:/certa-drdo/certa/Draft_TA.docx' 
    pdf_final_path ='E:/certa-drdo/certa/Draft_TA.pdf' 
    final_path='E:/certa-drdo/certa/'
    if os.path.isfile(pdf_final_path):
        with open(pdf_final_path, 'rb') as pdf:
            response = HttpResponse(pdf.read(),content_type='application/pdf')
            response['Content-Disposition'] = 'filename=some_file.pdf'
        return response
    elif os.path.isfile(doc_final_path):
        print('mmmmmmmmmmmmmm')
        pythoncom.CoInitialize()
        wdFormatPDF = 17
        # print(tempfile.gettempdir(),'temp')

        in_file = os.path.abspath(doc_final_path)
        # out_file = os.path.abspath('D:/cemilac/certa/defence/media/org1.pdf')

        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs('E:/certa-drdo/certa/Draft_TA.pdf', FileFormat=wdFormatPDF)
        print('nnnnnnnnnnn')
        doc.Close()
        word.Quit()
        with open(final_path+'Draft_TA.pdf', 'rb') as pdf:
            response = HttpResponse(pdf.read(),content_type='application/pdf')
            response['Content-Disposition'] = 'filename=some_file.pdf'
        return response
    else:
    
        idprefix=request.POST['idprefix']
        print(idprefix,'jjjjjjjjjjjj')

        curr_path = "/"+str(id)+ "/"+idprefix+"Annexure 7/"
        curr_path=curr_path.replace('/','\\')
        new_path = os.path.join(settings.MEDIA_ROOT + curr_path)
    
        # if os.path.isdir(new_path):
        #     with open(new_path+'Draft_TA.pdf', 'rb') as pdf:
        #         response = HttpResponse(pdf.read(),content_type='application/pdf')
        #         response['Content-Disposition'] = 'filename=some_file.pdf'
        #     return response
        # else:
        taa=TAapplicationmodel.objects.filter(user_id=id).first()

        # template = get_template('dealing officer/Draft TA pdf.html')
        target_file = StringIO()
        template = DocxTemplate("E:/certa-drdo/certa/dashboard/templates/dealing officer/template.docx")
        context= {
            'firmname':taa.firmname,
            'addr1':taa.addr1,
            'item_name':taa.item_name,
            'part_no':taa.part_no

        }
        html = template.render(context)
        doc_io = io.BytesIO() # create a file-like object
        template.save("Draft_TA.docx") # save data to file-like object

        
        new_path1 = 'E:\certa-drdo\certa\Draft_TA.docx'
        output_path = os.path.join(settings.MEDIA_ROOT) + '/89/result.pdf'
        # new_path=new_path.replace('\','//')

        taa=TAapplicationfiles.objects.filter(user_id=id,refid=idprefix,refpath='Annexure 4.13').first()
        aesurl=taa.filepath
        docurl = aesurl[:-4]
        print('aesview',aesurl)
        print('docurl',docurl)

        bufferSize = 64 * 1024
        passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
        encFileSize = stat(aesurl).st_size
        with open(aesurl, "rb") as fIn:
            with open(docurl, "wb") as fOut:
                pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)

        # curr_path = "/"+str(id)+ "/Annexure 4.13/PC/pc.docx.aes"
        # curr_path=curr_path.replace('/','\\')
        
        # new_path = os.path.join(settings.MEDIA_ROOT + curr_path)
        # templateDoc = Document(new_path1)
        templateDoc1 = Document(new_path1)
        templateDoc = Document(docurl)

        templateDoc1.add_page_break()
        
        for element in templateDoc.element.body:
            templateDoc1.element.body.append(element)

        templateDoc1.save(new_path1)
        print(request.user.id,'kkkkkkkk')
        messages.success(request, 'Draft_TA Successfully Prepared, Click again to view the file !')
        reg=TAapplicationmodel.objects.filter(file_in_id=str(request.user.id),file_in_name="TCS-DO")
        print('reggggggg',reg)
        return render(request, 'tcs do/receivedtyperecord.html',{'details':reg,'status':True})

    #     pisaStatus = pisa.CreatePDF(
    #     html, dest=response, link_callback=link_callback)
    #     if pisaStatus:
    #         return HttpResponse(response, content_type='application/pdf')
    # # if error then show some funy view
    #     if pisaStatus.err:
    #         return HttpResponse('We had some errors <pre>' + html + '</pre>')
    #     return response

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD","TCS-CE","TCS-Dealing Officer","TCS-TA Coordinator"])
def data_sheet(request,id):

    idprefix=request.POST['idprefix']
    print(idprefix,'jjjjjjjjjjjj')

    doc_final_path ='E:/certa-drdo/certa/TA_Datasheet.docx' 
    final_path ='E:/certa-drdo/certa/' 
    # finalpath=final_path.replace('/','\\')
    pdf_final_path ='E:/certa-drdo/certa/TA_Datasheet.pdf' 
    if os.path.isfile(pdf_final_path):
        with open(pdf_final_path, 'rb') as pdf:
            response = HttpResponse(pdf.read(),content_type='application/pdf')
            response['Content-Disposition'] = 'filename=some_file.pdf'
        return response
    elif os.path.isfile(doc_final_path):
        print('mmmmmmmmmmmmmm')
        pythoncom.CoInitialize()
        wdFormatPDF = 17
        # print(tempfile.gettempdir(),'temp')

        in_file = os.path.abspath(doc_final_path)
        # out_file = os.path.abspath('D:/cemilac/certa/defence/media/org1.pdf')

        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs('E:/certa-drdo/certa/TA_Datasheet.pdf', FileFormat=wdFormatPDF)
        print('nnnnnnnnnnn')
        doc.Close()
        word.Quit()
        with open(final_path+'TA_Datasheet.pdf', 'rb') as pdf:
            response = HttpResponse(pdf.read(),content_type='application/pdf')
            response['Content-Disposition'] = 'filename=some_file.pdf'
        return response
    else:
        curr_path = "/"+str(id)+ "/"+idprefix+"Annexure 6/"
        curr_path=curr_path.replace('/','\\')
        new_path = os.path.join(settings.MEDIA_ROOT + curr_path)
    
        # if os.path.isdir(new_path):
        #     with open(new_path+'TA Datasheet.docx', 'rb') as pdf:
        #         response = HttpResponse(pdf.read(),content_type='application/pdf')
        #         response['Content-Disposition'] = 'filename=some_file.pdf'
        #     return response
        # else:
        taa=TAapplicationmodel.objects.filter(user_id=id).first()

            # template = get_template('dealing officer/Draft TA pdf.html')
        target_file = StringIO()
        template = DocxTemplate("E:/certa-drdo/certa/dashboard/templates/dealing officer/DS template.docx")
        context= {
            'firmname':taa.firmname,
            'addr1':taa.addr1,
            'item_name':taa.item_name,
            'part_no':taa.part_no

        }
        html = template.render(context)
        doc_io = io.BytesIO() # create a file-like object
        template.save("TA_Datasheet.docx") # save data to file-like object

        
        new_path1 = 'E:\certa-drdo\certa\TA_Datasheet.docx'
        # output_path = os.path.join(settings.MEDIA_ROOT) + '/89/result.pdf'
        # new_path=new_path.replace('\','//')

        taa=TAapplicationfiles.objects.filter(user_id=id,refid=idprefix,refpath='Annexure 6').first()
        aesurl=taa.filepath
        docurl = aesurl[:-4]
        print('aesview',aesurl)
        print('docurl',docurl)

        bufferSize = 64 * 1024
        passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
        encFileSize = stat(aesurl).st_size
        with open(aesurl, "rb") as fIn:
            with open(docurl, "wb") as fOut:
                pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)

        
        templateDoc1 = Document(new_path1)
        templateDoc = Document(docurl)

        # templateDoc1.add_page_break()
        
        for element in templateDoc.element.body:
            templateDoc1.element.body.append(element)

        templateDoc1.save(new_path1)
        messages.success(request, 'Data_sheet Successfully Prepared, Click again to view the file !')

        reg=TAapplicationmodel.objects.filter(file_in_id=str(request.user.id))
        return render(request, 'tcs do/receivedtyperecord.html',{'details':reg,'status':True})

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD","TCS-CE","TCS-Dealing Officer","TCS-TA Coordinator"])
def addproforma(request,id):
    idprefix=request.POST['idprefix']
    print(idprefix,'kkkkkkkkkk')
    fc=TAapplicationmodel.objects.filter(user_id=id,idprefix=idprefix).first()
    print(fc.idprefix,'kkk')
    # tafil=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory="TAapplication",refid=fc.idprefix).first()

    curr_path = "/"+str(fc.user_id)+ fc.idprefix+"Annexure 3/Proforma_A/"
    curr_path=curr_path.replace('/','\\')
    new_path = os.path.join(settings.MEDIA_ROOT + curr_path)
    if os.path.isdir(new_path):
        with open(new_path+'Proforma_A.pdf', 'rb') as pdf:
            response = HttpResponse(pdf.read(),content_type='application/pdf')
            response['Content-Disposition'] = 'filename=some_file.pdf'
        return response
    else:
        print('sai',fc.user_id,fc.idprefix)
        form = proforma_A_form(request=fc.user_id,idpre=fc.idprefix)
        pro=proforma_A_model.objects.filter(user_id=fc.user_id,idprefix=idprefix).first()
        taa=TAapplicationmodel.objects.filter(user_id=fc.user_id,idprefix=idprefix).first()
        if pro:
            template = get_template('dealing officer/proformapdf.html')
            date_joined = datetime.now()
            formatted_datetime = date_joined.strftime("%Y-%m-%d")
            print(formatted_datetime,'dte')
            taf=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory='DAL_MDI',refid=fc.idprefix).first()
            dalurl=''
            if taf:
                aesurl=taf.filepath
                if taf.ext=='.pdf':
                    pdfurl = aesurl[:-4]
                    print('aesview',aesurl)
                    print('pdfview',pdfurl)
                    bufferSize = 64 * 1024
                    passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
                    encFileSize = stat(aesurl).st_size
                    with open(aesurl, "rb") as fIn:
                        with open(pdfurl, "wb") as fOut:
                            pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)
                    pdfpath = pdfurl[25:]
                    print(pdfpath,'pppppppppp')
                    curr_path=pdfpath
                    dalurl='http://127.0.0.1:8000/media'+curr_path
                    print(dalurl,'pppp11111pppppp')
            taf=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory='BOM',refid=fc.idprefix).first()
            bomurl=''
            if taf:
                aesurl=taf.filepath
                if taf.ext=='.pdf':
                    pdfurl = aesurl[:-4]
                    print('aesview',aesurl)
                    print('pdfview',pdfurl)
                    bufferSize = 64 * 1024
                    passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
                    encFileSize = stat(aesurl).st_size
                    with open(aesurl, "rb") as fIn:
                        with open(pdfurl, "wb") as fOut:
                            pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)
                    pdfpath = pdfurl[25:]
                    print(pdfpath,'pppppppppp')
                    curr_path=pdfpath
                    bomurl='http://127.0.0.1:8000/media'+curr_path
                    print(bomurl,'pppp11111pppppp')
            taf=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory='Tech_Spec',refid=fc.idprefix).first()
            techspecurl=''
            if taf:
                aesurl=taf.filepath
                if taf.ext=='.pdf':
                    pdfurl = aesurl[:-4]
                    print('aesview',aesurl)
                    print('pdfview',pdfurl)
                    bufferSize = 64 * 1024
                    passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
                    encFileSize = stat(aesurl).st_size
                    with open(aesurl, "rb") as fIn:
                        with open(pdfurl, "wb") as fOut:
                            pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)
                    pdfpath = pdfurl[25:]
                    print(pdfpath,'pppppppppp')
                    curr_path=pdfpath
                    techspecurl='http://127.0.0.1:8000/media'+curr_path
                    print(techspecurl,'pppp11111pppppp')
                    
            context= {
                'firmname':taa.firmname,
                'addr1':taa.addr1,
                'addr2':taa.addr2,
                'item_name':taa.item_name,
                'part_no':taa.part_no,
                'desc':taa.desc,
                'dal_mdi':taa.dal_mdi,
                'bom':taa.bom,
                'sop_acbs':taa.sop_acbs,
                'pc': taa.pc,
                'tre':taa.tre,
                'otheritems':taa.otheritems,
                'dalurl':dalurl,
                'bomurl':bomurl,
                'techspecurl':techspecurl,

                
                'ta': pro.ta,
                'techspec': pro.techspec,
                'qts': pro.qts,
                'qtr': pro.qtr,
                'cd': pro.cd,
                'photo': pro.photo,
                'feedback': pro.feedback,
                'req': pro.req,
                'cost': pro.cost,
                'quantity': pro.quantity,
                'pc': pro.pc,
                'tacomments':pro.tacomments,
                'datenow':formatted_datetime
                
            }

            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="report.pdf"'
            html = template.render(context)

            pisaStatus = pisa.CreatePDF(
            html,dest=response,link_callback=link_callback)
            if pisaStatus:
                return HttpResponse(response,content_type='application/pdf')
        # if error then show some funy view
            if pisaStatus.err:
                return HttpResponse('We had some errors <pre>' + html + '</pre>')
            return response
        else:
            print(form.errors)
            return render(request, 'dealing officer/proforma.html', {'form': form,'id':id,'idprefix':idprefix})

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer"])
def generateproformapdf(request):
    id=request.POST['id']
    idprefix=request.POST['idprefix']
    print('saiiiiiiiiiiiiiii',id)
    
    fc=TAapplicationmodel.objects.filter(user_id=id,idprefix=idprefix).first()
    print(fc.idprefix,'kkk')
    tafil=TAapplicationfiles.objects.filter(user_id=fc.user_id,filecategory="TAapplication",refid=fc.idprefix).first()

 
    # return render(request, 'dealing officer/proforma.html')

    if request.method=='POST':
        # firstname=request.POST['firstname']
        # lastname=request.POST['lastname']
        # country=request.POST['country']
        # subject=request.POST['subject']
        # reg=get_object_or_404(registration,id=id)
        
        user=User.objects.get(pk=fc.user_id)
        form = proforma_A_form(request.POST,request=fc.user_id,idpre=fc.idprefix)
        if form.is_valid():
            pro= form.save(commit=False)
            pro.user = user
            pro.idprefix=fc.idprefix
            pro.save()

            taapp_form=TAapplicationmodel.objects.filter(user_id=pro.user_id,idprefix=fc.idprefix).first()
            print("pro_form",taapp_form.id)

            get_taap_id=statusmodel.objects.filter(TAA_id=taapp_form.id).first()
            get_taap_id.status='Ready_for_CL'
            get_taap_id.Ready_for_CL=datetime.now()
            get_taap_id.save()
            print("status",get_taap_id)

            messages.success(request, 'Proforma_A Successfully Prepared !')
            return render(request, 'dealing officer/proforma.html')

        #     print('firstname',request.POST['firmname'])
        #     firmname=request.POST['firmname']
        #     template = get_template('dealing officer/proformapdf.html')
        #     context= {
        #         'desc':request.POST['desc'],
        #         'item_name':request.POST['item_name'],
        #         'part_no':request.POST['part_no'],
        #         'dal_mdi':request.POST['dal_mdi'],
        #         'bom':request.POST['bom'],
        #         'sop_acbs':request.POST['sop_acbs'],
        #         'otheritems':request.POST['otheritems'],
        #         'firmname':request.POST['firmname'],
        #         'addr1':request.POST['addr1'],
        #         'addr2':request.POST['addr2'],
        #         'ta': request.POST['ta'],
        #         'techspec': request.POST['techspec'],
        #         'qts': request.POST['qts'],
        #         'qtr': request.POST['qtr'],
        #         'cd': request.POST['cd'],
        #         'tre': request.POST['tre'],
        #         'photo': request.POST['photo'],
        #         'feedback': request.POST['feedback'],
        #         'req': request.POST['req'],
        #         'cost': request.POST['cost'],
        #         'quantity': request.POST['quantity'],
        #         'pc': request.POST['pc'],
        #         'tacomments':request.POST['tacomments']
        #     }
    
        #     response = HttpResponse(content_type='application/pdf')
        #     response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        #     html = template.render(context)

        #     pisaStatus = pisa.CreatePDF(
        #     html, dest=response, link_callback=link_callback)
        #     if pisaStatus:
        #         return HttpResponse(response, content_type='application/pdf')
        # # if error then show some funy view
        #     if pisaStatus.err:
        #         return HttpResponse('We had some errors <pre>' + html + '</pre>')
        #     return response
        else:
            print(form.errors)
    
@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD"])
def rowselect(request,id):
    form=commentsUploadForm
    print('if',id)
    idprefix=request.POST['idprefix']  
    print(idprefix,'idprefix')
    taf=TAapplicationfiles.objects.filter(user_id=id,refid=idprefix).order_by('refpath').first()
    get_refpath=TAapplicationfiles.objects.filter(user_id=id,refid=idprefix).values('refpath').order_by('refpath')
    idg=idgenerationmodel.objects.filter(user_id=id,idprefix=idprefix).first()
    print(get_refpath,'taff')
    
    for anex_name in get_refpath:
        anexture_name = anex_name['refpath']
        print(anexture_name,'taff')
        comments = commentsmodel(name=anexture_name,idprefix=idprefix,user_id=id)
        commentsdb=comments.save()
    Datadisp=commentsmodel.objects.filter(user_id=id,idprefix=idprefix).order_by('name')
    print(Datadisp,'Datadisp')
    # return render(request, 'applicant/view_all_doc.html',{'form':form,'details': taf,'idg':idg,'idprefix':idprefix})

    # taa=TAapplicationmodel.objects.filter(user_id=id).first()
    # taf=TAapplicationfiles.objects.filter(user_id=id).exclude(filecategory="TAapplication")
    # return render(request, 'dealing officer/detail view.html',{'taa':taa,'taf':taf,'id':id})
    return render(request, 'rd/comments_view_doc.html',{'form':form,'details': Datadisp})

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD"])
def addcomment(request):

    print('if',id)
    anexture_name=request.POST['name']
    comments=request.POST['comments']
    responsible=request.POST['responsible']
    status=request.POST['status']  
    idprefix=request.POST['idprefix']  
    print(idprefix,anexture_name,'idprefix')
    
    get_taap_id=commentsmodel.objects.filter(name=anexture_name,idprefix=idprefix).first()
    # get_taap_id.status='Send_to_RCMA_TAC'
    # get_taap_id.Send_to_RCMA_TAC=datetime.now()
    # get_taap_id.save()

    return render(request, 'rd/comments_view_doc.html')

@login_required(login_url=settings.LOGIN_URL)
@role_required(allowed_roles=["Dealing Officer","TA Coordinator","RD","TCS-GD"])
def pdfviewercopy(request,id):
    
    # curr_path = "/"+str(id)+ "/TAapplication/"
    # curr_path=curr_path.replace('/','\\')
    # new_path = os.path.join(settings.MEDIA_ROOT + curr_path)


    # with open(new_path+'TAapplication.pdf', 'rb') as pdf:
    #     response = HttpResponse(pdf.read(),content_type='application/pdf')
    #     response['Content-Disposition'] = 'filename=some_file.pdf'
    # return response
    taa=TAapplicationmodel.objects.filter(user_id=id).first()
    taf=TAapplicationfiles.objects.filter(user_id=id).exclude(filecategory="TAapplication")
    print('kkkkkkkkkkkkkkkkk')
    if request.POST:
        aesurl=request.POST['path']
        ext=request.POST['ext']
        tafnew=TAapplicationfiles.objects.filter(user_id=id,filepath=aesurl,ext=ext).first()
        fc=tafnew.comments
        print('aesview',aesurl)

        pdfurl=''
        docurl=''
        nameonly=''
        if ext=='.pdf':
            pdfurl = aesurl[:-3]+'pdf'
            print('aesview',aesurl)
            print('pdfview',pdfurl)
            bufferSize = 64 * 1024
            passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
            encFileSize = stat(aesurl).st_size
            with open(aesurl, "rb") as fIn:
                with open(pdfurl, "wb") as fOut:
                    pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)
            print(pdfurl,'pdfurl')
            pdfpath = pdfurl[25:]
            print(pdfpath)
            curr_path=pdfpath
            url='http://127.0.0.1:8000/media'+curr_path
            print(fc,'comments')
            return render(request, 'dealing officer/detail view.html',{'url':url,'id':id,'fc':fc,'taa':taa,'taf':taf,'path':aesurl})
        elif ext=='docx':
        # word to pdf 
            nameonly=aesurl[:-4]
            docurl = aesurl[:-4]+'.docx'
            print('aesview',aesurl)
            print('nameonly',nameonly)
            print('docurl',docurl)
            bufferSize = 64 * 1024
            passw = "#EX\xc8\xd5\xbfI{\xa2$\x05(\xd5\x18\xbf\xc0\x85)\x10nc\x94\x02)j\xdf\xcb\xc4\x94\x9d(\x9e"
            encFileSize = stat(aesurl).st_size
            with open(aesurl, "rb") as fIn:
                with open(docurl, "wb") as fOut:
                    pyAesCrypt.decryptStream(fIn, fOut, passw, bufferSize, encFileSize)

            pythoncom.CoInitialize()
            wdFormatPDF = 17
            in_file = os.path.abspath(docurl)
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(nameonly+'.pdf', FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()
            pdfurl=nameonly+'.pdf'
            print(pdfurl,'pdfurl')
            pdfpath = pdfurl[25:]
            print(pdfpath)
            curr_path=pdfpath
            url='http://127.0.0.1:8000/media'+curr_path
            print(fc,'comments')
            os.remove(docurl)
            return render(request, 'dealing officer/detail view.html',{'url':url,'id':id,'fc':fc,'taa':taa,'taf':taf,'path':aesurl})
       
            # with open(nameonly+'.pdf', 'rb') as pdf:
            #     response = HttpResponse(pdf.read(),content_type='application/pdf')
            #     response['Content-Disposition'] = 'filename=some_file.pdf'
            # return response
        # finally:
        #     os.remove(nameonly+'.pdf')
        #     os.remove(docurl)
    else:
        return render(request, 'dealing officer/detail view.html',{'id':id,'taa':taa,'taf':taf})
            
        # os.remove(pdfurl)
        # print('asdfasdfasdfasdfasdfds')

